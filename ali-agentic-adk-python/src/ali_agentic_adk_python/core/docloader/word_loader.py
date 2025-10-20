# Copyright (C) 2025 AIDC-AI
# This project incorporates components from the Open Source Software below.
# The original copyright notices and the licenses under which we received such components are set forth below for informational purposes.
#
# Open Source Software Licensed under the MIT License:
# --------------------------------------------------------------------
# 1. vscode-extension-updater-gitlab 3.0.1 https://www.npmjs.com/package/vscode-extension-updater-gitlab
# Copyright (c) Microsoft Corporation. All rights reserved.
# Copyright (c) 2015 David Owens II
# Copyright (c) Microsoft Corporation.
# Terms of the MIT:
# --------------------------------------------------------------------
# MIT License
#
# Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation files (the "Software"), to deal in the Software without restriction, including without limitation the rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and to permit persons to whom the Software is furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.

from __future__ import annotations

from io import BufferedReader, BytesIO
from pathlib import Path
from typing import Any, BinaryIO, Optional

from docx import Document as WordDocument

from ali_agentic_adk_python.core.docloader.base import BaseLoader
from ali_agentic_adk_python.core.indexes import Document


class WordDocLoader(BaseLoader):
    """Load DOCX files and convert paragraphs into ``Document`` objects."""

    def __init__(self, file_path: Optional[str] = None) -> None:
        self.file_path = file_path

    def load(self) -> list[Document]:
        if not self.file_path:
            raise ValueError("WordDocLoader requires `file_path` or metadata-driven loading.")
        return self._load_from_path(self.file_path)

    def fetch_content(self, document_meta: dict[str, Any]) -> list[Document]:
        metadata_hint = document_meta.get("metadata")

        file_path = document_meta.get("file_path") or document_meta.get("filePath")
        if file_path:
            documents = self._load_from_path(str(file_path))
        elif stream := (
            document_meta.get("input_stream")
            or document_meta.get("stream")
            or document_meta.get("binary_stream")
        ):
            documents = self._load_from_stream(stream)
        elif raw_bytes := document_meta.get("bytes"):
            documents = self._load_from_stream(BytesIO(raw_bytes))
        else:
            documents = super().fetch_content(document_meta)

        if metadata_hint:
            for doc in documents:
                doc.metadata.update(metadata_hint)
        return documents

    def _load_from_path(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        with path.open("rb") as fh:
            return self._extract_documents(fh, str(path))

    def _load_from_stream(
        self, stream: BinaryIO | BufferedReader | BytesIO
    ) -> list[Document]:
        return self._extract_documents(stream, None)

    def _extract_documents(self, stream: BinaryIO, source: Optional[str]) -> list[Document]:
        docx = WordDocument(stream)
        paragraphs = [para.text.strip() for para in docx.paragraphs if para.text.strip()]
        if not paragraphs:
            text = ""
        else:
            text = "\n\n".join(paragraphs)
        metadata = {"source": source or "stream"}
        return [Document(page_content=text, metadata=metadata)]
